#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LexControl - Índice de texto completo de los expedientes (PDF, DOCX, OCR)
========================================================================
Extrae el texto de todos los PDF y documentos Word (DOCX) del disco forense y lo deja en un índice SQLite FTS5.
Soporta OCR automático con Tesseract para PDFs escaneados (imágenes).

Es incremental: guarda el tamaño y la fecha de cada archivo ya procesado, así que
volver a correrlo sólo indexa lo nuevo o lo que cambió. Se puede interrumpir con
Ctrl-C y retomar donde iba.

    python3 indexar_pdfs.py              # indexa lo que falte
    python3 indexar_pdfs.py --rehacer    # rehace el índice desde cero
    python3 indexar_pdfs.py --limite 200 # sólo los primeros 200 (para probar)
"""
import argparse
import hashlib
import json
import os
import sqlite3
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import pytesseract
    from PIL import Image
    import io
except ImportError:
    pytesseract = None

fitz.TOOLS.mupdf_display_errors(False)

BASE_DIR = Path(__file__).resolve().parent
INDICE = BASE_DIR / "data" / "indice_texto.sqlite"

RAICES = [
    Path("/media/jaime/c11cad3b-6d38-462a-9c2e-49c33f1f6c18/Casos2023"),
    Path("/home/jaime/Descargas/Casos2023-Consolidados"),
]

MAX_PAGINAS = 60
MAX_CARACTERES = 400_000

# --- Índice semántico (embeddings) --------------------------------------------
# nomic-embed-text vía Ollama, 768 dimensiones. El modelo corre cargado con
# contexto fijo de 2048 tokens; pedir más (con "num_ctx" en la request o con
# /api/embed + truncate=True) igual responde 500 "input length exceeds the
# context length" -se verificó empíricamente (30-jul-2026) que Ollama no trunca
# solo. Por eso truncamos nosotros: 6000 caracteres funciona, 8000 falla.
OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434").rstrip("/")
MODELO_EMBEDDING = "nomic-embed-text"
MAX_CARACTERES_EMBEDDING = 6000


def abrir_indice():
    INDICE.parent.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(INDICE)
    con.execute("PRAGMA journal_mode=WAL")
    con.execute("""
        CREATE TABLE IF NOT EXISTS archivos (
            ruta      TEXT PRIMARY KEY,
            nombre    TEXT,
            carpeta   TEXT,
            tamano    INTEGER,
            modificado REAL,
            paginas   INTEGER,
            indexado  REAL
        )
    """)
    con.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS textos USING fts5(
            ruta UNINDEXED,
            nombre,
            carpeta,
            contenido,
            tokenize = "unicode61 remove_diacritics 2"
        )
    """)
    con.execute("""
        CREATE TABLE IF NOT EXISTS embeddings (
            ruta      TEXT PRIMARY KEY,
            modelo    TEXT,
            vector    BLOB,
            generado  REAL
        )
    """)
    # Hash SHA256 del contenido, para detectar el mismo archivo bajado dos veces
    # con nombre distinto (ej. "Resolución (1).pdf") -por ruta no se detecta,
    # porque cada descarga repetida es un path nuevo-. Se agrega con ALTER en vez
    # de ir en el CREATE TABLE de arriba porque el índice ya existe en producción
    # con miles de filas: así las bases viejas se actualizan solas al abrirlas.
    try:
        con.execute("ALTER TABLE archivos ADD COLUMN hash TEXT")
    except sqlite3.OperationalError:
        pass  # la columna ya existe
    con.execute("CREATE INDEX IF NOT EXISTS idx_archivos_hash ON archivos(hash)")
    return con


def calcular_hash(ruta):
    """SHA256 del contenido del archivo. None si no se puede leer -nunca lanza-."""
    try:
        h = hashlib.sha256()
        with open(ruta, "rb") as f:
            for bloque in iter(lambda: f.read(1024 * 1024), b""):
                h.update(bloque)
        return h.hexdigest()
    except Exception:
        return None


def calcular_hash_bytes(datos):
    return hashlib.sha256(datos).hexdigest() if datos else None


def embeber_texto(texto):
    """Pide a Ollama el vector semántico (768 floats, nomic-embed-text) de `texto`.
    Devuelve una lista de floats, o None si Ollama no está disponible o falla
    -nunca lanza-, para que el llamador pueda seguir sin romperse (igual criterio
    que _ollama_generar_json en servidor_local_lexcontrol.py)."""
    payload = {"model": MODELO_EMBEDDING, "prompt": texto[:MAX_CARACTERES_EMBEDDING]}
    try:
        req = urllib.request.Request(
            f"{OLLAMA_HOST}/api/embeddings",
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"}
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data.get("embedding")
    except Exception:
        return None


def guardar_embedding(con, ruta, vector):
    """Guarda `vector` (lista de floats) como BLOB float32 en la tabla `embeddings`."""
    arr = np.asarray(vector, dtype=np.float32)
    con.execute(
        "INSERT OR REPLACE INTO embeddings (ruta, modelo, vector, generado) VALUES (?, ?, ?, ?)",
        (str(ruta), MODELO_EMBEDDING, arr.tobytes(), time.time())
    )


def extraer_texto_docx(ruta):
    if not docx:
        return None, 0
    try:
        doc = docx.Document(ruta)
        parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parrafos.append(row_text)
        texto = "\n".join(parrafos)[:MAX_CARACTERES]
        return texto, 1
    except Exception as e:
        return None, 0


def extraer_texto_pdf(ruta):
    try:
        with fitz.open(ruta) as doc:
            total = doc.page_count
            partes = []
            ocr_necesario = False
            
            for i, pagina in enumerate(doc):
                if i >= MAX_PAGINAS:
                    break
                txt = pagina.get_text()
                if txt and len(txt.strip()) > 30:
                    partes.append(txt)
                else:
                    # Si no hay texto directo y tenemos pytesseract, intentamos OCR en la página
                    if pytesseract:
                        try:
                            pix = pagina.get_pixmap(dpi=150)
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            txt_ocr = pytesseract.image_to_string(img, lang="spa")
                            if txt_ocr and len(txt_ocr.strip()) > 20:
                                partes.append(f"[OCR] {txt_ocr}")
                        except Exception:
                            pass
                if sum(len(p) for p in partes) > MAX_CARACTERES:
                    break
            return "\n".join(partes)[:MAX_CARACTERES], total
    except Exception as e:
        return None, 0


def extraer_texto(ruta):
    p = Path(ruta)
    ext = p.suffix.lower()
    if ext == ".docx":
        return extraer_texto_docx(ruta)
    elif ext == ".pdf":
        return extraer_texto_pdf(ruta)
    elif ext in [".txt", ".md"]:
        try:
            with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
                return f.read(MAX_CARACTERES), 1
        except Exception:
            return None, 0
    return None, 0


def recolectar_archivos():
    raiz = next((r for r in RAICES if r.exists()), None)
    if raiz is None:
        print("❌ No se encuentra el disco de casos. Conéctalo y vuelve a intentar.")
        for r in RAICES:
            print(f"     {r}")
        sys.exit(1)
    print(f"📂 Recorriendo {raiz}")
    extensiones = {".pdf", ".docx", ".txt"}
    archivos = sorted(str(p) for p in raiz.rglob("*") if p.suffix.lower() in extensiones and p.is_file())
    return raiz, archivos


def main():
    parser = argparse.ArgumentParser(description="Índice FTS5 de expedientes (PDF, DOCX, OCR)")
    parser.add_argument("--rehacer", action="store_true", help="Borra el índice actual y lo recrea")
    parser.add_argument("--limite", type=int, default=0, help="Límite de archivos a procesar")
    args = parser.parse_args()

    if args.rehacer and INDICE.exists():
        print(f"⚠️  Borrando índice anterior: {INDICE}")
        INDICE.unlink()

    con = abrir_indice()
    raiz, lista_archivos = recolectar_archivos()

    ya = {
        ruta: (tam, mod)
        for ruta, tam, mod in con.execute("SELECT ruta, tamano, modificado FROM archivos")
    }

    pendientes = []
    for ruta in lista_archivos:
        try:
            st = os.stat(ruta)
        except OSError:
            continue
        previo = ya.get(ruta)
        if previo and previo[0] == st.st_size and abs(previo[1] - st.st_mtime) < 1:
            continue
        pendientes.append((ruta, st))

    if args.limite:
        pendientes = pendientes[: args.limite]

    print(f"🔎 {len(pendientes)} documentos por indexar ({len(lista_archivos) - len(pendientes)} ya al día).")
    if not pendientes:
        print("✅ El índice de contenido está 100% al día.")
        return

    inicio = time.time()
    hechos = fallidos = 0
    try:
        for n, (ruta, st) in enumerate(pendientes, 1):
            texto, paginas = extraer_texto(ruta)
            if not texto:
                fallidos += 1
                continue

            p = Path(ruta)
            carpeta = p.parent.name
            con.execute("DELETE FROM textos WHERE ruta = ?", (ruta,))
            con.execute(
                "INSERT INTO textos (ruta, nombre, carpeta, contenido) VALUES (?, ?, ?, ?)",
                (ruta, p.name, carpeta, texto),
            )
            con.execute(
                "INSERT OR REPLACE INTO archivos (ruta, nombre, carpeta, tamano, modificado, paginas, indexado, hash)"
                " VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (ruta, p.name, carpeta, st.st_size, st.st_mtime, paginas, time.time(), calcular_hash(ruta)),
            )
            hechos += 1

            if n % 100 == 0:
                con.commit()
                transcurrido = time.time() - inicio
                ritmo = n / transcurrido if transcurrido else 0
                faltan = (len(pendientes) - n) / ritmo if ritmo else 0
                print(f"   {n}/{len(pendientes)}  ({ritmo:.1f} doc/s, faltan ~{faltan/60:.1f} min)")
    except KeyboardInterrupt:
        print("\n⏸️  Interrumpido. Lo indexado queda guardado; vuelve a correrlo para continuar.")
    finally:
        con.commit()
        con.execute("INSERT INTO textos(textos) VALUES('optimize')")
        con.commit()
        con.close()

    total = time.time() - inicio
    print(f"\n✅ {hechos} documentos indexados con éxito en {total/60:.1f} min. {fallidos} sin texto extraíble.")
    print(f"   Índice local: {INDICE}  ({INDICE.stat().st_size/1048576:.1f} MB)")


if __name__ == "__main__":
    main()
